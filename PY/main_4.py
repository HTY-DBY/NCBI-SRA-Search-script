import os
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt, RGBColor

from PY.GSet import GSet
from PY.Other import load_json_my

# 设置测试标志和分隔符
Test_do = 0
split_str = '-->'

# 根据测试标志选择文件路径
Translated_json_path = GSet().Translated_json_save if Test_do == 0 else GSet().Translated_json_save_test

# 加载 JSON 数据
Translated_data = load_json_my(Translated_json_path)
new_dict = {}

# 初始化 new_dict 并填充数据
for index, item in enumerate(Translated_data.items()):
	search_key = item[0].split(split_str)[0]
	new_dict.setdefault(search_key, {}).update({item[1].get('Link'): item[1]})

import docx


def add_hyperlink(paragraph, url, text, color, underline):
	"""
	A function that places a hyperlink within a paragraph object.

	:param paragraph: The paragraph we are adding the hyperlink to.
	:param url: A string containing the required url
	:param text: The text displayed for the url
	:return: The hyperlink object
	"""

	# This gets access to the document.xml.rels file and gets a new relation id value
	part = paragraph.part
	r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

	# Create the w:hyperlink tag and add needed values
	hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
	hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

	# Create a w:r element
	new_run = docx.oxml.shared.OxmlElement('w:r')

	# Create a new w:rPr element
	rPr = docx.oxml.shared.OxmlElement('w:rPr')

	# Add color if it is given
	if not color is None:
		c = docx.oxml.shared.OxmlElement('w:color')
		c.set(docx.oxml.shared.qn('w:val'), color)
		rPr.append(c)

	# Remove underlining if it is requested
	if not underline:
		u = docx.oxml.shared.OxmlElement('w:u')
		u.set(docx.oxml.shared.qn('w:val'), 'none')
		rPr.append(u)

	# Join all the xml elements together add add the required text to the w:r element
	new_run.append(rPr)
	new_run.text = text
	hyperlink.append(new_run)

	paragraph._p.append(hyperlink)

	return hyperlink


def save_search_key_to_word(item):
	search_key = item[0]
	Data_ALL_in_search_key = item[1]

	try:
		doc = Document()
		style = doc.styles['Normal']
		style.font.name = 'Times New Roman'
		style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

		# 设置页边距
		section = doc.sections[0]
		section.top_margin = Inches(0.5)
		section.bottom_margin = Inches(0.5)
		section.left_margin = Inches(0.5)
		section.right_margin = Inches(0.5)

		# 填充文档内容
		for index, (key, entry) in enumerate(Data_ALL_in_search_key.items()):
			doc.add_heading(f'{index + 1}: {entry["Title"]}', level=1)
			temp = ' ---------------------------------------------------------'
			content_sections = [
				("Title", f"Title:{temp}", RGBColor(255, 0, 0)),
				("Title", entry["Title"], RGBColor(0, 0, 0)),
				("translated_Title", entry.get("translated_Title", ""), RGBColor(0, 0, 0)),
				("Link", f'Link:', RGBColor(255, 0, 0)),
				("Submitted", f"Submitted:{temp}", RGBColor(255, 0, 0)),
				("Submitted", entry["Submitted"], RGBColor(0, 0, 0)),
				("translated_Submitted", entry.get("translated_Submitted", ""), RGBColor(0, 0, 0)),
				("Design", f"Design:{temp}", RGBColor(255, 0, 0)),
				("Design", entry["Design"], RGBColor(0, 0, 0)),
				("translated_Design", entry.get("translated_Design", ""), RGBColor(0, 0, 0)),
				("Study", f"Study:{temp}", RGBColor(255, 0, 0)),
				("Study", entry["Study"], RGBColor(0, 0, 0)),
				("translated_Study", entry.get("translated_Study", ""), RGBColor(0, 0, 0)),
				("Library", f"Library:{temp}", RGBColor(255, 0, 0)),
				("Library", entry["Library"], RGBColor(0, 0, 0)),
				("translated_Library", entry.get("translated_Library", ""), RGBColor(0, 0, 0)),
				("BioData", f'BioData: {entry["BioData"]}', RGBColor(255, 0, 0)),
				("Separator", '====================================================', RGBColor(0, 0, 0))
			]

			# 写入每一段内容
			for section_key, text, color in content_sections:
				if text:  # Only process non-empty text
					p = doc.add_paragraph(text)
					run = p.runs[0]
					run.font.color.rgb = color
					p.paragraph_format.space_after = Pt(0)
				else:
					# If text is empty, add a placeholder or skip
					p = doc.add_paragraph("No content available")
					run = p.runs[0]
					run.font.color.rgb = RGBColor(0, 0, 0)  # Lighter color for "No content"
					p.paragraph_format.space_after = Pt(0)

				if section_key == "Link":

					p = doc.add_paragraph()
					# Ensure entry["Link"] is not None
					if entry.get("Link"):
						hyperlink = add_hyperlink(p, entry["Link"], entry["Link"], '#0e69c4', True)
					else:
						p.add_run("No Link Available").font.color.rgb = RGBColor(255, 0, 0)
					p.paragraph_format.space_after = Pt(0)

		output_file = os.path.join(GSet().Data_fin if Test_do == 0 else GSet().Data_temp, f'{search_key}.docx')
		doc.save(output_file)
		return search_key, output_file
	except Exception as e:
		print(f"Error saving search_key '{search_key}': {e}")
		return search_key, None


# 使用线程池并行处理每个 search_key 并显示进度
with ThreadPoolExecutor(max_workers=os.cpu_count()) as executor:
	futures = {executor.submit(save_search_key_to_word, item): item[0] for item in new_dict.items()}

	for i, future in enumerate(as_completed(futures), start=1):
		try:
			search_key, output_file = future.result()
			if output_file:
				print(f"进度 {i}/{len(new_dict)} - {search_key}\n已保存到 {output_file}")
			else:
				print(f"进度 {i}/{len(new_dict)} - {search_key} 文件保存失败")
		except Exception as e:
			print(f"Error processing search_key '{futures[future]}': {e}")
